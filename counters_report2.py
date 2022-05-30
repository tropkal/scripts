# This script works the same as counters_report.py, but instead of opening the output with notepad, this will use Microsoft Word to open it because I wanted to have some of
# the text centered and bolded.

from pysnmp.entity.rfc3413.oneliner import cmdgen
import os
import time
import sys
from docx import Document # pip install python-docx
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

PATH_TO_OUTFILE = '<path>'

# Instantiate Document() class
document = Document()

# IPs of your network printers and its corresponding description for easier management
IPs = {
  
    'ip1' : 'description1',
    'ip2' : 'description2',
    'ip3' : 'description3'
}

# You need to have the same number of IPs and printers (each printer has its own (probably) specific OIDs that you need to query for)
COPY_MACHINE_1 = {
  
    '1.3.6.1.4.1.18334.1.1.1.5.7.2.2.1.5.1.1' : 'copy counter black',
    '1.3.6.1.4.1.18334.1.1.1.5.7.2.2.1.5.2.1' : 'copy counter full color',
    '1.3.6.1.4.1.18334.1.1.1.5.7.2.2.1.5.3.1' : 'copy counter single color',
    '1.3.6.1.4.1.18334.1.1.1.5.7.2.2.1.7.1.1' : 'copy counter black large size',
    '1.3.6.1.4.1.18334.1.1.1.5.7.2.2.1.7.2.1' : 'copy counter full color large size',
    '1.3.6.1.4.1.18334.1.1.1.5.7.2.2.1.7.4.1' : 'copy counter single color large size',
  
    '1.3.6.1.4.1.18334.1.1.1.5.7.2.2.1.5.1.2' : 'print counter black',
    '1.3.6.1.4.1.18334.1.1.1.5.7.2.2.1.5.2.2' : 'print counter full color',
    '1.3.6.1.4.1.18334.1.1.1.5.7.2.2.1.5.4.2' : 'print counter 2 color', 
    '1.3.6.1.4.1.18334.1.1.1.5.7.2.2.1.7.1.2' : 'print counter black large size',
    '1.3.6.1.4.1.18334.1.1.1.5.7.2.2.1.7.2.2' : 'print counter full color large size',
    '1.3.6.1.4.1.18334.1.1.1.5.7.2.2.1.7.4.2' : 'print counter 2 color large size',

    '1.3.6.1.4.1.18334.1.1.1.5.7.2.1.5.0'     : 'scan counter total',
    '1.3.6.1.4.1.18334.1.1.1.5.7.2.3.1.6.1'   : 'scan counter large size',

    '1.3.6.1.2.1.43.11.1.1.9.1.3'             : 'yellow toner',
    '1.3.6.1.2.1.43.11.1.1.9.1.2'             : 'magenta toner',
    '1.3.6.1.2.1.43.11.1.1.9.1.1'             : 'cyan toner',
    '1.3.6.1.2.1.43.11.1.1.9.1.4'             : 'black toner'
}

COPY_MACHINE_2 = {
  
    '1.3.6.1.4.1.18334.1.1.1.5.7.2.2.1.5.1.1' : 'copy counter black',
    '1.3.6.1.4.1.18334.1.1.1.5.7.2.2.1.5.1.2' : 'print counter black',
    '1.3.6.1.4.1.18334.1.1.1.5.7.2.1.5.0'     : 'scan counter total',

    '1.3.6.1.2.1.43.11.1.1.9.1.1'             : 'black toner'
}

COPY_MACHINE_3 = {
  
    '1.3.6.1.4.1.18334.1.1.1.5.7.2.2.1.5.1.1' : 'copy counter black',
    '1.3.6.1.4.1.18334.1.1.1.5.7.2.2.1.5.1.2' : 'print counter black',
    '1.3.6.1.4.1.18334.1.1.1.5.7.2.1.5.0'     : 'scan counter total',
    '1.3.6.1.4.1.18334.1.1.1.5.7.2.3.1.6.1'   : 'scan counter large size',

    '1.3.6.1.2.1.43.11.1.1.9.1.1'             : 'black toner'
}

MACHINES = [ COPY_MACHINE_1, COPY_MACHINE_2, COPY_MACHINE_3 ]

def queryMachine(ip, counter):
    SNMP_RO_COMM = 'public'

    # Define a PySNMP CommunityData object named `auth`, by providing the SNMP community string
    auth = cmdgen.CommunityData(SNMP_RO_COMM)

    # Define the CommandGenerator, which will be used to send SNMP queries
    cmdGen = cmdgen.CommandGenerator()

    # Query a network device using the getCmd() function, providing the auth object, a UDP transport
    # our OID for SYSNAME, and don't lookup the OID in PySNMP's MIBs
    errorIndication, errorStatus, errorIndex, varBinds = cmdGen.getCmd(
        auth,
        cmdgen.UdpTransportTarget((ip, 161)),
        cmdgen.MibVariable(counter),
        lookupMib=False,
    )

    # Check if there was an error querying the device
    if errorIndication:
        sys.exit()

    # We only expect a single response from the host for sysName, but varBinds is an object
    # that we need to iterate over. It provides the OID and the value, both of which have a
    # prettyPrint() method so that you can get the actual string data
    
    for oid, val in varBinds:

        if val.prettyPrint() == 'No Such Instance currently exists at this OID':
            
            pass

        else:

            paragraph = document.add_paragraph()
            paragraph.add_run('[!] Counter: ' + MACHINES[i][counter] + '\n')

            if 'toner' in MACHINES[i][counter]:

                paragraph = document.add_paragraph()
                paragraph.add_run('[+] Value of counter: ' + str(val.prettyPrint()) + '%' + '\n')

            else:

                paragraph = document.add_paragraph()
                paragraph.add_run('[+] Value of counter: ' + '{:,}'.format(int(val.prettyPrint())) + '\n')

            document.add_paragraph('----------------------------------------------------------------------------------------------------------------------\n')

            return
# Start of document
paragraph = document.add_paragraph()
run = paragraph.add_run('[*] Today\'s date is: ' + time.asctime(time.localtime(time.time())) + '\n')

# Center and bold the date
paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
run.bold = True
document.add_paragraph('----------------------------------------------------------------------------------------------------------------------\n')

for i, ip in enumerate(IPs.keys()):

    paragraph = document.add_paragraph()
    run = paragraph.add_run('[!] IP is: ' + ip + ' (' + IPs[ip] + ')' + '\n')
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run.bold = True

    for counter in MACHINES[i].keys():
        queryMachine(ip, counter)

# Save the document after you're done editing it.
document.save(PATH_TO_OUTFILE)

# Use Word to automatically open the report. Add it to your %PATH%, otherwise this won't work.
os.system(f"WINWORD.exe {PATH_TO_OUTFILE}")
